from docx import Document
import json

# List of Word documents to process
WORD_LIST = [
    "raw/W1_Tarifas transferencias Extranjero.docx",
    "raw/W2_Ficha Tecnica.docx",
    "raw/W3_Catalogo de productos de activo vigentes.docx",
    "raw/W4_2023_12_Posicionamiento_Environment-1.docx",
    "raw/W5_2023_12_Posicionamiento_Environment-2.docx",
]


def extract_text_and_tables_from_docx(docx_path):
    """
    Extract text sections and tables from a .docx file.
    Handles irregular tables by normalizing row lengths.
    """
    document = Document(docx_path)
    sections = []
    tables = []

    current_section = {"title": "", "content": ""}

    # Process paragraphs
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            # Save the current section if it has content
            if current_section["title"] or current_section["content"]:
                sections.append(current_section)
                current_section = {"title": "", "content": ""}
            current_section["title"] = paragraph.text.strip()
        else:
            current_section["content"] += paragraph.text.strip() + " "

    # Append the last section if it exists
    if current_section["title"] or current_section["content"]:
        sections.append(current_section)

    # Process tables
    for table in document.tables:
        table_data = []
        max_cols = max(len(row.cells) for row in table.rows)
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            # Normalize irregular tables by filling missing columns with empty strings
            while len(row_data) < max_cols:
                row_data.append("")
            table_data.append(row_data)
        tables.append(table_data)

    return {"sections": sections, "tables": tables}


def process_all_docx_files(docx_files):
    """
    Process all .docx files in the provided list and return extracted data.
    """
    all_data = {}
    for docx_file in docx_files:
        try:
            all_data[docx_file] = extract_text_and_tables_from_docx(docx_file)
        except Exception as e:
            print(f"Error processing file {docx_file}: {e}")
    return all_data


def save_to_json(data, output_path):
    """
    Save extracted data to a JSON file.
    """
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


# Process the .docx files and save the results
extracted_data = process_all_docx_files(WORD_LIST)
save_to_json(extracted_data, "extracted_word_data.json")

print("Extraction complete. Data saved to 'extracted_word_data.json'.")
