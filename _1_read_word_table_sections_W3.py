from docx import Document
import json


def parse_docx_to_json(docx_path):
    """
    Parse a .docx file into a JSON format for chatbot.
    """
    document = Document(docx_path)
    parsed_data = {"sections": [], "tables": []}

    current_section = {"title": "", "content": ""}
    section_stack = []

    # Extract text sections
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue  # Skip empty paragraphs

        # Check for headings to structure sections
        if paragraph.style.name.startswith("Heading"):
            level = int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
            title = text

            # Save the current section
            if current_section["title"] or current_section["content"]:
                parsed_data["sections"].append(current_section)
                current_section = {"title": "", "content": ""}

            # Manage hierarchical structure with a stack
            while len(section_stack) >= level:
                section_stack.pop()
            section_stack.append(title)

            # Create a full hierarchical title
            full_title = " > ".join(section_stack)
            current_section["title"] = full_title
        else:
            # Append content to the current section
            current_section["content"] += text + " "

    # Add the last section
    if current_section["title"] or current_section["content"]:
        parsed_data["sections"].append(current_section)

    # Extract tables
    for table in document.tables:
        table_data = []
        table_title = ""

        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]
            if row_idx == 0:
                # Assume the first row contains the table title or headers
                table_title = row_data[0] if row_data else "Tabla sin título"
            else:
                table_data.append(row_data)

        parsed_data["tables"].append({
            "title": table_title,
            "data": table_data
        })

    return parsed_data


def save_to_json(data, output_path):
    """
    Save parsed data into a JSON file.
    """
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


# File path for input and output
INPUT_FILE = "raw/W3_Catalogo de productos de activo vigentes.docx"
OUTPUT_FILE = "processed/W3_Catalogo de productos de activo vigentes.json"

# Parse the file and save the output
parsed_data = parse_docx_to_json(INPUT_FILE)
save_to_json(parsed_data, OUTPUT_FILE)

print(f"Processing complete. JSON saved to {OUTPUT_FILE}")
