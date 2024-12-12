from docx import Document
import json
import re
from transformers import pipeline

# Load summarization pipeline
summarizer = pipeline("summarization", model="facebook/bart-large-cnn", tokenizer="facebook/bart-large-cnn")

# Input file path
FILE_PATH = "raw/W1_Tarifas transferencias Extranjero.docx"


def generate_title(content, max_length=10):
    """
    Generate a title for the given paragraph content using a language model.
    """
    try:
        summary = summarizer(content, max_length=max_length, min_length=3, do_sample=False)
        return summary[0]['summary_text']
    except Exception as e:
        print(f"Error generating title for content: {content[:50]}... -> {e}")
        return "Título generado automáticamente"


def remove_uniform_columns(data):
    """
    Remove columns where all elements are identical across rows.
    """
    if not data or len(data) < 2:  # No data or no rows to compare
        return data

    headers = data[0]
    rows = data[1:]
    num_columns = len(headers)

    # Check for uniform columns
    to_remove = []
    for col_idx in range(num_columns):
        column_values = [row[col_idx] for row in rows if col_idx < len(row)]
        if all(value == column_values[0] for value in column_values):
            to_remove.append(col_idx)

    # Remove uniform columns
    filtered_data = []
    for row in data:
        filtered_row = [cell for idx, cell in enumerate(row) if idx not in to_remove]
        filtered_data.append(filtered_row)

    return filtered_data


def parse_docx_to_json(docx_path):
    """
    Parse a .docx file into structured sections, tables, and standalone notes.
    """
    document = Document(docx_path)
    parsed_data = {"sections": [], "tables": []}

    current_section = {"title": "", "content": ""}

    # Parse sections and standalone notes
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue  # Skip empty paragraphs

        # Match standalone notes like "A.-", "B.-", "F.-"
        if re.match(r"^[A-Z]\.-", text):
            parsed_data["sections"].append({"title": "", "content": text})
            continue

        # Identify sections by paragraph style
        if paragraph.style.name.startswith("Heading"):
            # Save the current section if it has content
            if current_section["title"] or current_section["content"]:
                parsed_data["sections"].append(current_section)
                current_section = {"title": "", "content": ""}
            current_section["title"] = text
        else:
            current_section["content"] += text + " "

    # Append the last section if it exists
    if current_section["title"] or current_section["content"]:
        parsed_data["sections"].append(current_section)

    # Generate titles for sections without a title
    for section in parsed_data["sections"]:
        if not section["title"] and section["content"]:
            section["title"] = generate_title(section["content"])

    # Parse tables
    for table_idx, table in enumerate(document.tables, start=1):
        table_data = []
        table_title = ""

        for row_idx, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells]

            if row_idx == 0:
                # Use the first cell as the table title
                table_title = row_data[0] if row_data else f"Table {table_idx}"
                continue  # Skip the title row when processing data

            table_data.append(row_data)

        # Remove uniform columns
        table_data = remove_uniform_columns(table_data)

        # Append the table with its title
        parsed_data["tables"].append({
            "title": table_title,
            "data": table_data
        })

    # Add the last three static tables manually
    last_tables = [
        {
            "title": "PAÍSES DE LA UE (28)",
            "data": [
                ["País", "Cod. ISO", "Cod.B.E."],
                ["Alemania", "DE", "004"],
                ["Austria", "AT", "038"],
                ["Bélgica", "BE", "017"],
                ["Bulgaria", "BG", "068"],
                ["Chipre", "CY", "600"],
                ["Croacia", "HR", "092"],
                ["Dinamarca", "DK", "008"],
                ["Eslovaquia", "SK", "063"],
                ["Eslovenia", "SI", "091"],
                ["España", "ES", "011"],
                ["Estonia", "EE", "053"],
                ["Finlandia", "FI", "032"],
                ["Francia", "FR", "001"],
                ["Grecia", "GR", "009"],
                ["Hungría", "HU", "064"],
                ["Irlanda", "IE", "007"],
                ["Italia", "IT", "005"],
                ["Letonia", "LV", "054"],
                ["Lituania", "LT", "055"],
                ["Luxemburgo", "LU", "018"],
                ["Malta", "MT", "046"],
                ["Países Bajos", "NL", "003"],
                ["Polonia", "PL", "060"],
                ["Portugal", "PT", "010"],
                ["Reino Unido", "GB", "006"],
                ["Rep. Checa", "CZ", "061"],
                ["Rumanía", "RO", "066"],
                ["Suecia", "SE", "030"]
            ]
        },
        {
            "title": "PAÍSES EN CONVENIO CON LA UE",
            "data": [
                ["País", "Cod. ISO", "Cod.B.E."],
                ["Noruega", "NO", "028"],
                ["Islandia", "IS", "024"],
                ["Liechtenstein", "LI", "037"]
            ]
        },
        {
            "title": "OTROS CONCEPTOS (C)",
            "data": [
                ["%", "Normal", "Urgente", "SWIFT"],
                ["----", "10 €", "12 €", "----"],
                ["Datos insuficientes o incorrectos", "----", "15 €", "----"],
                ["Transferencias Urgentes Divisa", "----", "15 €", "----"],
                ["Investigación, Modificación o Aclaración", "----", "15 €", "----"]
            ]
        }
    ]

    for table in last_tables:
        table["data"] = remove_uniform_columns(table["data"])
        parsed_data["tables"].append(table)

    return parsed_data


def save_to_json(data, output_path):
    """
    Save parsed data to a JSON file.
    """
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


# Process the .docx file
parsed_data = parse_docx_to_json(FILE_PATH)

# Save the output to JSON
OUTPUT_PATH = "processed/W1_Tarifas_transferencias.json"
save_to_json(parsed_data, OUTPUT_PATH)

print(f"Processing complete. Data saved to {OUTPUT_PATH}")
